"""
Cover Letter Generator Module
==============================
Stage 6 of the pipeline.

Responsibilities:
- Uses LLM to write a tailored, professional cover letter.
- Based on: candidate resume + target job description.
- Outputs: CoverLetter model + DOCX file.

Design: The cover letter is always written fresh (no template filling).
        The LLM has full creative freedom within the constraints defined
        in the system prompt to ensure quality and personalization.
"""

import logging
import uuid
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

from core.config import settings
from core.llm_client import LLMClient
from models.schemas import CoverLetter, OptimizedResume, ParsedJobDescription
from prompts.templates import COVER_LETTER_SYSTEM, COVER_LETTER_USER

logger = logging.getLogger(__name__)


class CoverLetterGenerator:
    """
    Generates tailored cover letters using LLM + formats as DOCX.

    Usage:
        generator = CoverLetterGenerator(llm_client)
        letter, path = await generator.generate(resume, parsed_jd)
    """

    def __init__(self, llm_client: LLMClient):
        self.llm = llm_client
        self.output_dir = Path(settings.OUTPUT_DIR)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    async def generate(
        self,
        resume: OptimizedResume,
        parsed_jd: ParsedJobDescription,
    ) -> tuple[CoverLetter, str]:
        """
        Generate cover letter content via LLM and save as DOCX.
        Returns: (CoverLetter model, path to DOCX file)
        """
        logger.info(f"Generating cover letter for: {parsed_jd.job_title}")

        # Build user prompt with resume + JD context
        recent_role = ""
        if resume.work_experience:
            exp = resume.work_experience[0]
            recent_role = f"{exp.title} at {exp.company}"

        user_prompt = COVER_LETTER_USER.format(
            name=resume.name,
            summary=resume.summary or "Experienced professional",
            skills=", ".join(resume.skills[:15]),
            recent_role=recent_role or "Previous role",
            achievements="\n".join(f"- {a}" for a in resume.achievements[:5]),
            job_title=parsed_jd.job_title,
            company_name=parsed_jd.company_name or "your company",
            responsibilities="\n".join(f"- {r}" for r in parsed_jd.key_responsibilities[:6]),
            required_skills=", ".join(parsed_jd.required_skills[:10]),
        )

        raw_json = await self.llm.generate_json(
            system_prompt=COVER_LETTER_SYSTEM,
            user_prompt=user_prompt,
        )

        cover_letter = self._json_to_model(raw_json, parsed_jd)
        docx_path = self._generate_docx(cover_letter, resume.name)

        return cover_letter, docx_path

    # -----------------------------------------------------------------------
    # JSON → CoverLetter Model
    # -----------------------------------------------------------------------
    def _json_to_model(self, data: dict, parsed_jd: ParsedJobDescription) -> CoverLetter:
        """Convert LLM JSON output to CoverLetter model."""
        try:
            opening = data.get("opening_paragraph", "")
            body1 = data.get("body_paragraph_1", "")
            body2 = data.get("body_paragraph_2", "")
            closing = data.get("closing_paragraph", "")

            # Assemble full text if not provided by LLM
            full_text = data.get("full_text", "")
            if not full_text:
                recipient = data.get("recipient_name", "Hiring Manager")
                full_text = (
                    f"Dear {recipient},\n\n"
                    f"{opening}\n\n"
                    f"{body1}\n\n"
                    f"{body2}\n\n"
                    f"{closing}\n\n"
                    f"Sincerely,"
                )

            return CoverLetter(
                recipient_name=data.get("recipient_name", "Hiring Manager"),
                company_name=data.get("company_name") or parsed_jd.company_name or "the company",
                job_title=data.get("job_title") or parsed_jd.job_title,
                opening_paragraph=opening,
                body_paragraph_1=body1,
                body_paragraph_2=body2,
                closing_paragraph=closing,
                full_text=full_text,
            )
        except Exception as e:
            logger.error(f"Error parsing cover letter JSON: {e}")
            # Return a minimal cover letter on error
            return CoverLetter(
                company_name=parsed_jd.company_name or "the company",
                job_title=parsed_jd.job_title,
                opening_paragraph="I am writing to express my strong interest in this position.",
                body_paragraph_1="",
                body_paragraph_2="",
                closing_paragraph="I look forward to discussing this opportunity further.",
                full_text="Cover letter generation encountered an issue. Please regenerate.",
            )

    # -----------------------------------------------------------------------
    # DOCX Cover Letter Generation
    # -----------------------------------------------------------------------
    def _generate_docx(self, letter: CoverLetter, candidate_name: str) -> str:
        """Generate a professional DOCX cover letter."""
        doc = Document()

        # Page setup
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

        # Default font
        from docx.oxml.ns import qn
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Date
        from datetime import date
        date_para = doc.add_paragraph()
        date_para.paragraph_format.space_after = Pt(12)
        date_run = date_para.add_run(date.today().strftime("%B %d, %Y"))
        date_run.font.size = Pt(11)

        # Salutation
        salutation = doc.add_paragraph()
        salutation.paragraph_format.space_after = Pt(12)
        salutation.add_run(f"Dear {letter.recipient_name or 'Hiring Manager'},")

        # Opening paragraph
        if letter.opening_paragraph:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            p.add_run(letter.opening_paragraph)

        # Body paragraphs
        if letter.body_paragraph_1:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            p.add_run(letter.body_paragraph_1)

        if letter.body_paragraph_2:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            p.add_run(letter.body_paragraph_2)

        # Closing
        if letter.closing_paragraph:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(24)
            p.add_run(letter.closing_paragraph)

        # Sign-off
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph()
        sign = doc.add_paragraph()
        sign_run = sign.add_run(candidate_name)
        sign_run.font.bold = True

        # Save
        safe_name = "".join(c if c.isalnum() else "_" for c in candidate_name.lower())[:20]
        filename = f"cover_letter_{safe_name}_{uuid.uuid4().hex[:8]}.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        logger.info(f"Cover letter saved: {filepath}")
        return str(filepath)
