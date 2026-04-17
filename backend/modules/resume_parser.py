"""
Resume Parser Module
====================
Stage 1 of the pipeline.

Responsibilities:
1. Extract raw text from uploaded PDF or DOCX files.
2. Send extracted text to LLM for structured parsing.
3. Return a validated ParsedResume object.

Design: File extraction is deterministic (no LLM needed).
        Structuring is LLM-powered for maximum flexibility.
"""

import logging
from io import BytesIO
from pathlib import Path

import pdfplumber  # Best PDF text extractor (handles columns, tables)
from docx import Document

from core.llm_client import LLMClient
from models.schemas import ParsedResume, ContactInfo, WorkExperience, ProjectExperience, Education
from prompts.templates import RESUME_PARSER_SYSTEM, RESUME_PARSER_USER

logger = logging.getLogger(__name__)


class ResumeParser:
    """
    Parses resume files (PDF/DOCX) into structured ParsedResume objects.

    Usage:
        parser = ResumeParser(llm_client)
        parsed = await parser.parse(file_bytes, filename="resume.pdf")
    """

    def __init__(self, llm_client: LLMClient):
        self.llm = llm_client

    # -----------------------------------------------------------------------
    # Public API
    # -----------------------------------------------------------------------
    async def parse(self, file_bytes: bytes, filename: str) -> ParsedResume:
        """Full pipeline: file → text → LLM → ParsedResume."""
        logger.info(f"Parsing resume: {filename}")

        # Step 1: Extract raw text
        raw_text = self._extract_text(file_bytes, filename)
        if not raw_text or len(raw_text.strip()) < 50:
            raise ValueError("Could not extract readable text from the resume file.")

        logger.info(f"Extracted {len(raw_text)} characters from resume")

        # Step 2: Parse with LLM
        return await self.parse_from_text(raw_text)

    async def parse_from_text(self, resume_text: str) -> ParsedResume:
        """Parse a resume from raw text string (used in both file upload and text input flows)."""
        user_prompt = RESUME_PARSER_USER.format(resume_text=resume_text[:8000])  # Truncate for token limits

        raw_json = await self.llm.generate_json(
            system_prompt=RESUME_PARSER_SYSTEM,
            user_prompt=user_prompt,
        )

        return self._json_to_model(raw_json)

    # -----------------------------------------------------------------------
    # Text Extraction (deterministic, no LLM)
    # -----------------------------------------------------------------------
    def _extract_text(self, file_bytes: bytes, filename: str) -> str:
        """Route to the appropriate extractor based on file extension."""
        ext = Path(filename).suffix.lower()

        if ext == ".pdf":
            return self._extract_from_pdf(file_bytes)
        elif ext in (".docx", ".doc"):
            return self._extract_from_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {ext}. Upload PDF or DOCX.")

    def _extract_from_pdf(self, file_bytes: bytes) -> str:
        """
        Extract text from PDF using pdfplumber.
        pdfplumber handles multi-column layouts better than PyPDF2.
        """
        text_parts = []
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text with layout preservation
                page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if page_text:
                    text_parts.append(page_text)
                    logger.debug(f"PDF page {page_num}: extracted {len(page_text)} chars")

        return "\n\n".join(text_parts)

    def _extract_from_docx(self, file_bytes: bytes) -> str:
        """
        Extract text from DOCX using python-docx.
        Preserves paragraph structure which helps the LLM parse sections.
        """
        doc = Document(BytesIO(file_bytes))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables (skills sometimes stored in tables)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        return "\n".join(paragraphs)

    # -----------------------------------------------------------------------
    # JSON → Pydantic Model Conversion
    # -----------------------------------------------------------------------
    def _json_to_model(self, data: dict) -> ParsedResume:
        """
        Safely convert LLM JSON output to a validated ParsedResume.
        Handles missing fields gracefully.
        """
        try:
            # Parse nested contact
            contact_data = data.get("contact", {}) or {}
            contact = ContactInfo(
                email=contact_data.get("email"),
                phone=contact_data.get("phone"),
                linkedin=contact_data.get("linkedin"),
                github=contact_data.get("github"),
                portfolio=contact_data.get("portfolio"),
                location=contact_data.get("location"),
            )

            # Parse work experience list
            work_exp = []
            for exp in data.get("work_experience", []):
                if isinstance(exp, dict):
                    work_exp.append(WorkExperience(
                        company=exp.get("company", "Unknown Company"),
                        title=exp.get("title", "Unknown Title"),
                        start_date=exp.get("start_date"),
                        end_date=exp.get("end_date"),
                        location=exp.get("location"),
                        bullets=[b for b in exp.get("bullets", []) if b],
                    ))

            # Parse projects list
            projects = []
            for proj in data.get("projects", []):
                if isinstance(proj, dict):
                    projects.append(ProjectExperience(
                        name=proj.get("name", "Project"),
                        description=proj.get("description"),
                        technologies=proj.get("technologies", []),
                        bullets=[b for b in proj.get("bullets", []) if b],
                        url=proj.get("url"),
                    ))

            # Parse education list
            education = []
            for edu in data.get("education", []):
                if isinstance(edu, dict):
                    education.append(Education(
                        institution=edu.get("institution", "Unknown Institution"),
                        degree=edu.get("degree", "Degree"),
                        field=edu.get("field"),
                        graduation_year=edu.get("graduation_year"),
                        gpa=edu.get("gpa"),
                        honors=edu.get("honors"),
                    ))

            return ParsedResume(
                name=data.get("name", ""),
                contact=contact,
                summary=data.get("summary"),
                skills=data.get("skills", []),
                work_experience=work_exp,
                projects=projects,
                education=education,
                achievements=data.get("achievements", []),
                certifications=data.get("certifications", []),
            )

        except Exception as e:
            logger.error(f"Error converting LLM JSON to ParsedResume: {e}", exc_info=True)
            # Return a minimal valid object rather than crashing
            return ParsedResume(name=data.get("name", "Unknown"))
