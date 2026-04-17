"""
DOCX Resume Generator
======================
Stage 5 of the pipeline.

Responsibilities:
- Takes an OptimizedResume model.
- Generates a clean, ATS-friendly DOCX file using python-docx.
- Follows strict ATS formatting rules (no tables, no columns, simple layout).
- Returns the path to the generated file.

ATS Formatting Rules Applied:
- Single-column layout only
- Standard section headings (no icons, no colors that break parsing)
- Bullet points using Word's native list style
- Standard fonts (Calibri 11pt)
- No headers/footers
- No text boxes or frames
"""

import logging
import uuid
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

from core.config import settings
from models.schemas import OptimizedResume

logger = logging.getLogger(__name__)


class DocxResumeGenerator:
    """
    Generates ATS-optimized DOCX resumes from OptimizedResume models.

    Design: python-docx for maximum compatibility.
    No external dependencies or templates needed.
    """

    # Color palette (subtle, professional)
    COLOR_NAME = RGBColor(0x1F, 0x49, 0x7D)     # Dark blue for name
    COLOR_HEADING = RGBColor(0x1F, 0x49, 0x7D)  # Same blue for section headings
    COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)    # Pure black for body text
    COLOR_GRAY = RGBColor(0x55, 0x55, 0x55)     # Gray for dates/secondary info

    FONT_NAME = "Calibri"
    FONT_SIZE_NAME = 18        # Candidate name
    FONT_SIZE_CONTACT = 10    # Contact details line
    FONT_SIZE_HEADING = 12    # Section headings
    FONT_SIZE_BODY = 11       # Body text

    def __init__(self):
        output_dir = Path(settings.OUTPUT_DIR)
        output_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir = output_dir

    def generate(self, resume: OptimizedResume, job_title: str = "") -> str:
        """
        Generate an ATS-friendly DOCX resume.
        Returns the absolute path to the generated file.
        """
        logger.info(f"Generating DOCX resume for: {resume.name}")
        doc = Document()

        # -----------------------------------------------------------------------
        # Page Setup: US Letter, 1-inch margins, no header/footer
        # -----------------------------------------------------------------------
        self._setup_page(doc)

        # -----------------------------------------------------------------------
        # Section 1: Name (centered, prominent)
        # -----------------------------------------------------------------------
        self._add_name(doc, resume.name)

        # -----------------------------------------------------------------------
        # Section 2: Contact Details (single line, centered)
        # -----------------------------------------------------------------------
        self._add_contact(doc, resume.contact)

        # -----------------------------------------------------------------------
        # Section 3: Professional Summary
        # -----------------------------------------------------------------------
        if resume.summary:
            self._add_section_heading(doc, "PROFESSIONAL SUMMARY")
            self._add_paragraph(doc, resume.summary)

        # -----------------------------------------------------------------------
        # Section 4: Core Skills (grouped, pipe-separated for ATS)
        # -----------------------------------------------------------------------
        if resume.skills:
            self._add_section_heading(doc, "CORE SKILLS")
            self._add_skills(doc, resume.skills)

        # -----------------------------------------------------------------------
        # Section 5: Work Experience
        # -----------------------------------------------------------------------
        if resume.work_experience:
            self._add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
            for exp in resume.work_experience:
                self._add_experience_entry(doc, exp)

        # -----------------------------------------------------------------------
        # Section 6: Project Experience
        # -----------------------------------------------------------------------
        if resume.projects:
            self._add_section_heading(doc, "PROJECT EXPERIENCE")
            for proj in resume.projects:
                self._add_project_entry(doc, proj)

        # -----------------------------------------------------------------------
        # Section 7: Education
        # -----------------------------------------------------------------------
        if resume.education:
            self._add_section_heading(doc, "EDUCATION")
            for edu in resume.education:
                self._add_education_entry(doc, edu)

        # -----------------------------------------------------------------------
        # Section 8: Achievements & Certifications
        # -----------------------------------------------------------------------
        all_achievements = list(resume.achievements) + list(resume.certifications)
        if all_achievements:
            self._add_section_heading(doc, "ACHIEVEMENTS & CERTIFICATIONS")
            for achievement in all_achievements:
                self._add_bullet(doc, achievement)

        # -----------------------------------------------------------------------
        # Save to output directory
        # -----------------------------------------------------------------------
        filename = f"resume_{self._safe_name(resume.name)}_{uuid.uuid4().hex[:8]}.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        logger.info(f"Resume saved: {filepath}")
        return str(filepath)

    # -----------------------------------------------------------------------
    # Page & Document Setup
    # -----------------------------------------------------------------------
    def _setup_page(self, doc: Document):
        """Configure page margins and default paragraph spacing."""
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

        # Set default style for the document
        style = doc.styles["Normal"]
        style.font.name = self.FONT_NAME
        style.font.size = Pt(self.FONT_SIZE_BODY)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)

    # -----------------------------------------------------------------------
    # Header: Name & Contact
    # -----------------------------------------------------------------------
    def _add_name(self, doc: Document, name: str):
        """Add candidate name as the document header."""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(name.upper())
        run.font.name = self.FONT_NAME
        run.font.size = Pt(self.FONT_SIZE_NAME)
        run.font.bold = True
        run.font.color.rgb = self.COLOR_NAME

    def _add_contact(self, doc: Document, contact):
        """Add contact details as a single centered line."""
        parts = []
        if contact.phone:
            parts.append(contact.phone)
        if contact.email:
            parts.append(contact.email)
        if contact.linkedin:
            # Simplify URL for readability
            linkedin = contact.linkedin.replace("https://", "").replace("http://", "")
            parts.append(linkedin)
        if contact.github:
            github = contact.github.replace("https://", "").replace("http://", "")
            parts.append(github)
        if contact.portfolio:
            portfolio = contact.portfolio.replace("https://", "").replace("http://", "")
            parts.append(portfolio)
        if contact.location:
            parts.append(contact.location)

        contact_line = "  |  ".join(parts)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(contact_line)
        run.font.name = self.FONT_NAME
        run.font.size = Pt(self.FONT_SIZE_CONTACT)
        run.font.color.rgb = self.COLOR_GRAY

    # -----------------------------------------------------------------------
    # Section Headings (with horizontal rule below)
    # -----------------------------------------------------------------------
    def _add_section_heading(self, doc: Document, heading_text: str):
        """Add a bold section heading with a bottom border line (ATS-safe divider)."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(2)

        run = para.add_run(heading_text)
        run.font.name = self.FONT_NAME
        run.font.size = Pt(self.FONT_SIZE_HEADING)
        run.font.bold = True
        run.font.color.rgb = self.COLOR_HEADING

        # Add bottom border to simulate a horizontal rule (ATS-safe)
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F497D")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # -----------------------------------------------------------------------
    # Content Adders
    # -----------------------------------------------------------------------
    def _add_paragraph(self, doc: Document, text: str, italic: bool = False):
        """Add a regular body paragraph."""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = self.FONT_NAME
        run.font.size = Pt(self.FONT_SIZE_BODY)
        run.font.italic = italic
        run.font.color.rgb = self.COLOR_BLACK

    def _add_skills(self, doc: Document, skills: list):
        """
        Add skills in groups of 4-5 per line, separated by | for ATS parsing.
        This is more ATS-friendly than bullet points for skills.
        """
        chunk_size = 5
        chunks = [skills[i:i+chunk_size] for i in range(0, len(skills), chunk_size)]
        for chunk in chunks:
            line = "  •  ".join(chunk)
            self._add_paragraph(doc, line)

    def _add_experience_entry(self, doc: Document, exp):
        """Add a work experience entry with title, company, dates, and bullets."""
        # Title | Company line (bold)
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(1)

        title_run = para.add_run(f"{exp.title}")
        title_run.font.bold = True
        title_run.font.name = self.FONT_NAME
        title_run.font.size = Pt(self.FONT_SIZE_BODY)

        company_run = para.add_run(f"  –  {exp.company}")
        company_run.font.name = self.FONT_NAME
        company_run.font.size = Pt(self.FONT_SIZE_BODY)

        # Date line (right-aligned)
        date_parts = []
        if exp.start_date:
            date_parts.append(exp.start_date)
        if exp.end_date:
            date_parts.append(exp.end_date)
        if exp.location:
            date_parts.append(exp.location)

        if date_parts:
            date_para = doc.add_paragraph()
            date_para.paragraph_format.space_after = Pt(2)
            date_run = date_para.add_run(" | ".join(date_parts))
            date_run.font.name = self.FONT_NAME
            date_run.font.size = Pt(10)
            date_run.font.italic = True
            date_run.font.color.rgb = self.COLOR_GRAY

        # Bullet points
        for bullet in exp.bullets:
            self._add_bullet(doc, bullet)

    def _add_project_entry(self, doc: Document, proj):
        """Add a project entry with name, technologies, and bullets."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(1)

        # Project name (bold)
        name_run = para.add_run(proj.name)
        name_run.font.bold = True
        name_run.font.name = self.FONT_NAME
        name_run.font.size = Pt(self.FONT_SIZE_BODY)

        # Technologies
        if proj.technologies:
            tech_run = para.add_run(f"  |  {', '.join(proj.technologies)}")
            tech_run.font.name = self.FONT_NAME
            tech_run.font.size = Pt(self.FONT_SIZE_BODY)
            tech_run.font.italic = True

        # URL if available
        if proj.url:
            url_para = doc.add_paragraph()
            url_para.paragraph_format.space_after = Pt(1)
            url_run = url_para.add_run(proj.url)
            url_run.font.name = self.FONT_NAME
            url_run.font.size = Pt(10)
            url_run.font.color.rgb = self.COLOR_GRAY

        # Description
        if proj.description:
            self._add_paragraph(doc, proj.description)

        # Bullets
        for bullet in proj.bullets:
            self._add_bullet(doc, bullet)

    def _add_education_entry(self, doc: Document, edu):
        """Add an education entry."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(1)

        # Degree and institution
        degree_text = edu.degree
        if edu.field:
            degree_text += f" in {edu.field}"

        degree_run = para.add_run(f"{degree_text}")
        degree_run.font.bold = True
        degree_run.font.name = self.FONT_NAME
        degree_run.font.size = Pt(self.FONT_SIZE_BODY)

        inst_run = para.add_run(f"  –  {edu.institution}")
        inst_run.font.name = self.FONT_NAME
        inst_run.font.size = Pt(self.FONT_SIZE_BODY)

        # Details line: year, GPA, honors
        details = []
        if edu.graduation_year:
            details.append(edu.graduation_year)
        if edu.gpa:
            details.append(f"GPA: {edu.gpa}")
        if edu.honors:
            details.append(edu.honors)

        if details:
            detail_para = doc.add_paragraph()
            detail_para.paragraph_format.space_after = Pt(2)
            detail_run = detail_para.add_run(" | ".join(details))
            detail_run.font.name = self.FONT_NAME
            detail_run.font.size = Pt(10)
            detail_run.font.italic = True
            detail_run.font.color.rgb = self.COLOR_GRAY

    def _add_bullet(self, doc: Document, text: str):
        """Add an ATS-safe bullet point using Word's List Bullet style."""
        try:
            para = doc.add_paragraph(style="List Bullet")
        except KeyError:
            # Fallback if style not available
            para = doc.add_paragraph()
            para.add_run("• ")

        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.left_indent = Inches(0.25)
        run = para.add_run(text)
        run.font.name = self.FONT_NAME
        run.font.size = Pt(self.FONT_SIZE_BODY)
        run.font.color.rgb = self.COLOR_BLACK

    # -----------------------------------------------------------------------
    # Utilities
    # -----------------------------------------------------------------------
    @staticmethod
    def _safe_name(name: str) -> str:
        """Convert name to filesystem-safe string."""
        return "".join(c if c.isalnum() else "_" for c in name.lower())[:20]
